"""
GBSkillEngine 文档解析服务

支持PDF和DOCX文档的文本提取、表格提取、智能分块
PDF使用PyMuPDF(文本/渲染) + pdfplumber(表格)双引擎
"""
import os
import re
import logging
import tempfile
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class SectionChunk:
    """文档分块 - 按章节切分的文档片段"""
    section_number: str          # "4" / "4.1" / "附录A"
    section_title: str           # "技术要求"
    content: str                 # 该段的完整文本
    level: int                   # 章节层级
    page_range: Tuple[int, int]  # (start_page, end_page)
    tables: List[Dict[str, Any]] = field(default_factory=list)  # 该段包含的表格
    char_range: Tuple[int, int] = (0, 0)  # 在全文中的(start, end)位置


@dataclass
class ParsedDocument:
    """解析后的文档结构"""
    text: str  # 全文文本
    title: str = ""  # 文档标题
    sections: List[Dict[str, Any]] = field(default_factory=list)  # 章节结构
    tables: List[Dict[str, Any]] = field(default_factory=list)  # 表格数据
    metadata: Dict[str, Any] = field(default_factory=dict)  # 元数据
    chunks: List[SectionChunk] = field(default_factory=list)  # 智能分块
    page_images: Dict[int, str] = field(default_factory=dict)  # 页码→图片路径


class DocumentParser:
    """文档解析器"""
    
    def __init__(self):
        self._fitz = None  # PyMuPDF
        self._docx = None  # python-docx
        self._pdfplumber = None  # pdfplumber
        self._pil = None  # Pillow
    
    def _get_fitz(self):
        """延迟加载PyMuPDF"""
        if self._fitz is None:
            try:
                import fitz
                self._fitz = fitz
            except ImportError:
                logger.warning("PyMuPDF未安装，PDF解析将不可用")
                raise ImportError("PyMuPDF未安装，请安装: pip install PyMuPDF")
        return self._fitz
    
    def _get_docx(self):
        """延迟加载python-docx"""
        if self._docx is None:
            try:
                import docx
                self._docx = docx
            except ImportError:
                logger.warning("python-docx未安装，DOCX解析将不可用")
                raise ImportError("python-docx未安装，请安装: pip install python-docx")
        return self._docx
    
    def _get_pdfplumber(self):
        """延迟加载pdfplumber"""
        if self._pdfplumber is None:
            try:
                import pdfplumber
                self._pdfplumber = pdfplumber
            except ImportError:
                logger.warning("pdfplumber未安装，PDF表格提取将不可用")
                return None
        return self._pdfplumber
    
    def _get_pillow(self):
        """延迟加载Pillow"""
        if self._pil is None:
            try:
                from PIL import Image
                self._pil = Image
            except ImportError:
                logger.warning("Pillow未安装，PDF页面渲染将不可用")
                return None
        return self._pil
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析文档
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            ParsedDocument: 解析结果
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文档文件不存在: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext in [".doc", ".docx"]:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """解析PDF文档 - 使用PyMuPDF(文本) + pdfplumber(表格)双引擎"""
        fitz = self._get_fitz()
        plumber = self._get_pdfplumber()
        
        text_parts = []
        all_tables = []
        page_char_offsets = []  # 每页文本在全文中的起始偏移
        
        try:
            doc = fitz.open(file_path)
            plumber_pdf = None
            if plumber:
                try:
                    plumber_pdf = plumber.open(file_path)
                except Exception as e:
                    logger.warning(f"pdfplumber打开PDF失败，仅使用PyMuPDF: {e}")
            
            current_offset = 0
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # PyMuPDF: 提取文本
                page_text = page.get_text("text")
                page_char_offsets.append(current_offset)
                text_parts.append(page_text)
                current_offset += len(page_text) + 1  # +1 for \n join
                
                # pdfplumber: 提取该页表格
                if plumber_pdf:
                    page_tables = self._extract_pdf_tables_pdfplumber(
                        plumber_pdf.pages[page_num], page_num + 1
                    )
                    all_tables.extend(page_tables)
                
                # 文本中的表格标记(作为补充)
                text_table_markers = self._extract_tables_from_text(page_text, page_num + 1)
                # 将标记信息合并到pdfplumber表格(如果有对应的)
                self._merge_table_markers(all_tables, text_table_markers)
            
            if plumber_pdf:
                plumber_pdf.close()
            doc.close()
            
            full_text = "\n".join(text_parts)
            
            # 提取标题
            title = self._extract_title(full_text)
            
            # 提取章节结构(增强版，带位置信息)
            sections = self._extract_sections_with_positions(full_text)
            
            # 构建智能分块
            chunks = self._build_section_chunks(
                full_text, sections, all_tables, page_char_offsets
            )
            
            return ParsedDocument(
                text=full_text,
                title=title,
                sections=sections,
                tables=all_tables,
                metadata={
                    "file_type": "pdf",
                    "page_count": len(text_parts),
                    "table_count": len(all_tables),
                    "chunk_count": len(chunks),
                    "has_pdfplumber": plumber is not None,
                },
                chunks=chunks,
            )
            
        except Exception as e:
            logger.error(f"PDF解析失败: {str(e)}")
            raise RuntimeError(f"PDF解析失败: {str(e)}")
    
    def _extract_pdf_tables_pdfplumber(
        self, plumber_page, page_num: int
    ) -> List[Dict[str, Any]]:
        """使用pdfplumber提取单页PDF表格"""
        tables = []
        
        try:
            extracted = plumber_page.extract_tables(
                table_settings={
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "snap_tolerance": 5,
                    "join_tolerance": 5,
                    "min_words_vertical": 1,
                    "min_words_horizontal": 1,
                }
            )
            
            if not extracted:
                # 降级: 尝试更宽松的策略
                extracted = plumber_page.extract_tables(
                    table_settings={
                        "vertical_strategy": "text",
                        "horizontal_strategy": "text",
                        "snap_tolerance": 8,
                        "join_tolerance": 8,
                    }
                )
            
            for table_idx, table_data in enumerate(extracted or []):
                if not table_data or len(table_data) < 2:
                    continue
                
                # 清洗数据: 去除None和空白
                cleaned_rows = []
                for row in table_data:
                    cleaned_row = [
                        (cell.strip() if cell else "") for cell in row
                    ]
                    # 跳过全空行
                    if any(c for c in cleaned_row):
                        cleaned_rows.append(cleaned_row)
                
                if len(cleaned_rows) < 2:
                    continue
                
                headers = cleaned_rows[0]
                rows = cleaned_rows[1:]
                
                # 尝试将数值字符串转为数字
                converted_rows = []
                for row in rows:
                    converted = []
                    for cell in row:
                        converted.append(self._try_parse_number(cell))
                    converted_rows.append(converted)
                
                tables.append({
                    "table_id": f"table_p{page_num}_{table_idx + 1}",
                    "headers": headers,
                    "rows": converted_rows,
                    "row_count": len(converted_rows),
                    "col_count": len(headers),
                    "page": page_num,
                    "extraction_method": "pdfplumber",
                })
                
        except Exception as e:
            logger.debug(f"pdfplumber提取第{page_num}页表格失败: {e}")
        
        return tables
    
    @staticmethod
    def _try_parse_number(value: str):
        """尝试将字符串转为数字"""
        if not value or not isinstance(value, str):
            return value
        cleaned = value.strip().replace(" ", "")
        if not cleaned:
            return value
        try:
            if "." in cleaned:
                return float(cleaned)
            return int(cleaned)
        except (ValueError, TypeError):
            return value
    
    def _merge_table_markers(
        self,
        tables: List[Dict[str, Any]],
        markers: List[Dict[str, Any]],
    ):
        """将文本中的表格标记信息(title)合并到pdfplumber提取的表格"""
        for marker in markers:
            marker_page = marker.get("page")
            marker_title = marker.get("title", "")
            
            # 查找同一页的pdfplumber表格，添加title
            matched = False
            for table in tables:
                if (table.get("page") == marker_page
                        and not table.get("title")
                        and table.get("extraction_method") == "pdfplumber"):
                    table["title"] = marker_title
                    matched = True
                    break
            
            # 如果没有匹配的pdfplumber表格，保留标记作为纯文本表格标识
            if not matched:
                existing_ids = {t["table_id"] for t in tables}
                if marker.get("table_id") not in existing_ids:
                    tables.append(marker)
    
    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """解析DOCX文档"""
        docx = self._get_docx()
        
        text_parts = []
        tables = []
        
        try:
            doc = docx.Document(file_path)
            
            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格
            for table_idx, table in enumerate(doc.tables):
                table_data = self._extract_docx_table(table, table_idx + 1)
                if table_data:
                    tables.append(table_data)
            
            full_text = "\n".join(text_parts)
            
            # 提取标题
            title = self._extract_title(full_text)
            
            # 提取章节结构
            sections = self._extract_sections_with_positions(full_text)
            
            # 构建智能分块
            chunks = self._build_section_chunks(full_text, sections, tables, [])
            
            return ParsedDocument(
                text=full_text,
                title=title,
                sections=sections,
                tables=tables,
                metadata={
                    "file_type": "docx",
                    "paragraph_count": len(text_parts),
                    "table_count": len(tables),
                    "chunk_count": len(chunks),
                },
                chunks=chunks,
            )
            
        except Exception as e:
            logger.error(f"DOCX解析失败: {str(e)}")
            raise RuntimeError(f"DOCX解析失败: {str(e)}")
    
    def _extract_docx_table(self, table, table_idx: int) -> Optional[Dict[str, Any]]:
        """从DOCX提取表格数据"""
        try:
            rows = []
            headers = []
            
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                
                if row_idx == 0:
                    headers = row_data
                else:
                    # 尝试数值转换
                    converted = [self._try_parse_number(cell) for cell in row_data]
                    rows.append(converted)
            
            if not headers and not rows:
                return None
            
            return {
                "table_id": f"table_{table_idx}",
                "headers": headers,
                "rows": rows,
                "row_count": len(rows),
                "col_count": len(headers) if headers else (len(rows[0]) if rows else 0),
                "extraction_method": "docx",
            }
        except Exception as e:
            logger.warning(f"表格提取失败: {str(e)}")
            return None
    
    def _extract_tables_from_text(self, text: str, page_num: int) -> List[Dict[str, Any]]:
        """从文本中提取表格标记（表标题识别）"""
        tables = []
        
        # 查找表格标记（如"表1"、"表 2"等）
        table_pattern = r"表\s*(\d+)[：:\s]*([\u4e00-\u9fa5a-zA-Z0-9\s]*)"
        matches = re.finditer(table_pattern, text)
        
        for match in matches:
            table_num = match.group(1)
            table_title = match.group(2).strip() if match.group(2) else ""
            
            tables.append({
                "table_id": f"table_{table_num}",
                "title": table_title,
                "page": page_num,
                "raw_position": match.start(),
            })
        
        return tables
    
    def _extract_title(self, text: str) -> str:
        """提取文档标题"""
        lines = text.strip().split("\n")
        
        for line in lines[:10]:  # 检查前10行
            line = line.strip()
            # 标题通常包含"GB"或是较长的中文描述
            if line and (
                "GB" in line.upper() or 
                re.match(r"^[\u4e00-\u9fa5]{4,}", line)
            ):
                return line[:100]  # 限制长度
        
        return lines[0][:100] if lines else ""
    
    def _extract_sections_with_positions(self, text: str) -> List[Dict[str, Any]]:
        """提取章节结构（增强版，带位置信息）"""
        sections = []
        
        # 匹配章节标题模式（如"1 范围"、"4.2 技术要求"等）
        section_pattern = r"^(\d+(?:\.\d+)*)\s+([\u4e00-\u9fa5a-zA-Z\s]+)"
        
        for match in re.finditer(section_pattern, text, re.MULTILINE):
            section_num = match.group(1)
            section_title = match.group(2).strip()
            level = section_num.count(".") + 1
            
            sections.append({
                "number": section_num,
                "title": section_title,
                "level": level,
                "start_pos": match.start(),
            })
        
        # 检测附录
        appendix_pattern = r"^(附录\s*[A-Z])[：:\s]*([\u4e00-\u9fa5a-zA-Z\s]*)"
        for match in re.finditer(appendix_pattern, text, re.MULTILINE):
            sections.append({
                "number": match.group(1).replace(" ", ""),
                "title": match.group(2).strip() if match.group(2) else "",
                "level": 1,
                "start_pos": match.start(),
            })
        
        # 按位置排序
        sections.sort(key=lambda s: s.get("start_pos", 0))
        
        return sections
    
    def _build_section_chunks(
        self,
        full_text: str,
        sections: List[Dict[str, Any]],
        tables: List[Dict[str, Any]],
        page_char_offsets: List[int],
    ) -> List[SectionChunk]:
        """
        构建智能分块
        
        按一级章节(level=1)切分文本，超长段按二级章节再切分。
        每个chunk关联该段内的表格。
        """
        if not sections:
            # 没有章节结构，整个文档作为一个chunk
            if full_text.strip():
                return [SectionChunk(
                    section_number="0",
                    section_title="全文",
                    content=full_text,
                    level=0,
                    page_range=(1, max(1, len(page_char_offsets))),
                    tables=tables,
                    char_range=(0, len(full_text)),
                )]
            return []
        
        # 找出所有一级章节的位置
        level1_sections = [s for s in sections if s.get("level") == 1]
        
        if not level1_sections:
            level1_sections = sections[:1]
        
        chunks = []
        text_len = len(full_text)
        
        for i, sec in enumerate(level1_sections):
            start_pos = sec.get("start_pos", 0)
            
            # 结束位置: 下一个一级章节的开始，或文档末尾
            if i + 1 < len(level1_sections):
                end_pos = level1_sections[i + 1].get("start_pos", text_len)
            else:
                end_pos = text_len
            
            chunk_text = full_text[start_pos:end_pos]
            
            # 计算页码范围
            page_range = self._calc_page_range(
                start_pos, end_pos, page_char_offsets
            )
            
            # 查找该范围内的表格
            chunk_tables = self._find_tables_in_range(
                tables, start_pos, end_pos, page_range
            )
            
            # 如果chunk过长(>6000字符)，按二级章节再切分
            if len(chunk_text) > 6000:
                sub_chunks = self._split_large_chunk(
                    sec, chunk_text, start_pos, end_pos,
                    sections, tables, page_char_offsets
                )
                chunks.extend(sub_chunks)
            else:
                chunks.append(SectionChunk(
                    section_number=sec["number"],
                    section_title=sec["title"],
                    content=chunk_text,
                    level=sec.get("level", 1),
                    page_range=page_range,
                    tables=chunk_tables,
                    char_range=(start_pos, end_pos),
                ))
        
        # 处理第一个章节之前的前言内容
        first_start = level1_sections[0].get("start_pos", 0) if level1_sections else 0
        if first_start > 100:  # 有足够的前言内容
            preamble_text = full_text[:first_start]
            preamble_tables = self._find_tables_in_range(
                tables, 0, first_start,
                self._calc_page_range(0, first_start, page_char_offsets)
            )
            chunks.insert(0, SectionChunk(
                section_number="0",
                section_title="前言",
                content=preamble_text,
                level=0,
                page_range=self._calc_page_range(0, first_start, page_char_offsets),
                tables=preamble_tables,
                char_range=(0, first_start),
            ))
        
        return chunks
    
    def _split_large_chunk(
        self,
        parent_section: Dict[str, Any],
        chunk_text: str,
        abs_start: int,
        abs_end: int,
        all_sections: List[Dict[str, Any]],
        all_tables: List[Dict[str, Any]],
        page_char_offsets: List[int],
    ) -> List[SectionChunk]:
        """将过长的chunk按二级章节再切分"""
        parent_num = parent_section["number"]
        
        # 找该一级章节下的二级子章节
        sub_sections = [
            s for s in all_sections
            if (s.get("level", 1) == 2
                and s["number"].startswith(parent_num + ".")
                and abs_start <= s.get("start_pos", 0) < abs_end)
        ]
        
        if not sub_sections:
            # 没有二级子章节，保持原chunk
            page_range = self._calc_page_range(abs_start, abs_end, page_char_offsets)
            return [SectionChunk(
                section_number=parent_num,
                section_title=parent_section["title"],
                content=chunk_text,
                level=parent_section.get("level", 1),
                page_range=page_range,
                tables=self._find_tables_in_range(
                    all_tables, abs_start, abs_end, page_range
                ),
                char_range=(abs_start, abs_end),
            )]
        
        sub_chunks = []
        
        # 父章节头部到第一个子章节之间的内容
        first_sub_start = sub_sections[0].get("start_pos", abs_start)
        if first_sub_start > abs_start:
            head_text = chunk_text[:first_sub_start - abs_start]
            if head_text.strip():
                pr = self._calc_page_range(abs_start, first_sub_start, page_char_offsets)
                sub_chunks.append(SectionChunk(
                    section_number=parent_num,
                    section_title=parent_section["title"],
                    content=head_text,
                    level=parent_section.get("level", 1),
                    page_range=pr,
                    tables=self._find_tables_in_range(
                        all_tables, abs_start, first_sub_start, pr
                    ),
                    char_range=(abs_start, first_sub_start),
                ))
        
        for j, sub_sec in enumerate(sub_sections):
            sub_start = sub_sec.get("start_pos", abs_start)
            if j + 1 < len(sub_sections):
                sub_end = sub_sections[j + 1].get("start_pos", abs_end)
            else:
                sub_end = abs_end
            
            sub_text = chunk_text[sub_start - abs_start:sub_end - abs_start]
            pr = self._calc_page_range(sub_start, sub_end, page_char_offsets)
            
            sub_chunks.append(SectionChunk(
                section_number=sub_sec["number"],
                section_title=sub_sec["title"],
                content=sub_text,
                level=sub_sec.get("level", 2),
                page_range=pr,
                tables=self._find_tables_in_range(all_tables, sub_start, sub_end, pr),
                char_range=(sub_start, sub_end),
            ))
        
        return sub_chunks
    
    @staticmethod
    def _calc_page_range(
        start_pos: int, end_pos: int, page_char_offsets: List[int]
    ) -> Tuple[int, int]:
        """根据字符位置计算页码范围"""
        if not page_char_offsets:
            return (1, 1)
        
        start_page = 1
        end_page = len(page_char_offsets)
        
        for i, offset in enumerate(page_char_offsets):
            if offset <= start_pos:
                start_page = i + 1
            if offset <= end_pos:
                end_page = i + 1
        
        return (start_page, end_page)
    
    @staticmethod
    def _find_tables_in_range(
        tables: List[Dict[str, Any]],
        start_pos: int,
        end_pos: int,
        page_range: Tuple[int, int],
    ) -> List[Dict[str, Any]]:
        """查找在给定范围内的表格"""
        matched = []
        for table in tables:
            # 按页码匹配
            table_page = table.get("page")
            if table_page and page_range[0] <= table_page <= page_range[1]:
                matched.append(table)
                continue
            # 按字符位置匹配
            raw_pos = table.get("raw_position")
            if raw_pos is not None and start_pos <= raw_pos < end_pos:
                matched.append(table)
        return matched
    
    def render_page_to_image(
        self, file_path: str, page_num: int, output_dir: Optional[str] = None
    ) -> Optional[str]:
        """
        渲染PDF指定页面为JPEG图片
        
        Args:
            file_path: PDF文件路径
            page_num: 页码(从1开始)
            output_dir: 输出目录，默认使用临时目录
            
        Returns:
            图片文件路径，失败返回None
        """
        fitz = self._get_fitz()
        pil_image = self._get_pillow()
        if not pil_image:
            return None
        
        try:
            doc = fitz.open(file_path)
            if page_num < 1 or page_num > len(doc):
                doc.close()
                return None
            
            page = doc[page_num - 1]
            
            # 渲染为pixmap (DPI=150, 平衡质量和大小)
            mat = fitz.Matrix(150 / 72, 150 / 72)
            pix = page.get_pixmap(matrix=mat)
            
            # 转为Pillow Image
            img = pil_image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # 限制最大宽度2048px
            if img.width > 2048:
                ratio = 2048 / img.width
                new_size = (2048, int(img.height * ratio))
                img = img.resize(new_size, pil_image.LANCZOS)
            
            # 保存为JPEG
            if output_dir is None:
                output_dir = tempfile.mkdtemp(prefix="gbskill_pages_")
            os.makedirs(output_dir, exist_ok=True)
            
            img_path = os.path.join(output_dir, f"page_{page_num}.jpg")
            img.save(img_path, "JPEG", quality=85)
            
            doc.close()
            return img_path
            
        except Exception as e:
            logger.warning(f"渲染PDF第{page_num}页失败: {e}")
            return None
    
    def render_pages_to_images(
        self, file_path: str, page_nums: List[int], output_dir: Optional[str] = None
    ) -> Dict[int, str]:
        """
        批量渲染PDF页面为JPEG图片
        
        Returns:
            Dict[page_num, image_path]
        """
        if output_dir is None:
            output_dir = tempfile.mkdtemp(prefix="gbskill_pages_")
        
        result = {}
        for pn in page_nums:
            path = self.render_page_to_image(file_path, pn, output_dir)
            if path:
                result[pn] = path
        return result
    
    def extract_dimension_tables(self, parsed_doc: ParsedDocument) -> Dict[str, Any]:
        """
        从解析后的文档中提取尺寸规格表
        
        专门用于提取管材类国标中的DN-外径-壁厚对应表
        """
        dimension_tables = {}
        
        # 从表格中提取
        for table in parsed_doc.tables:
            headers = table.get("headers", [])
            rows = table.get("rows", [])
            
            if not headers or not rows:
                continue
            
            # 检测是否是尺寸表
            header_text = " ".join(str(h) for h in headers).lower()
            if any(kw in header_text for kw in ["dn", "外径", "壁厚", "公称"]):
                dimension_tables[table.get("table_id", "unknown")] = {
                    "headers": headers,
                    "data": rows,
                    "type": "dimension"
                }
        
        return dimension_tables


# 全局解析器实例
document_parser = DocumentParser()


def parse_standard_document(file_path: str) -> ParsedDocument:
    """
    解析国标文档
    
    Args:
        file_path: 文档路径
        
    Returns:
        ParsedDocument: 解析结果
    """
    return document_parser.parse(file_path)
